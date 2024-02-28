
import os
import pandas as pd
import pdfplumber
from docx import Document
from openpyxl import Workbook

# Define the directory containing PDF files
pdf_directory = r'C:\Users\hp\Downloads\New folder'

# Function to extract data from PDF files
def extract_data_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        pages = pdf.pages
        data = []
        for page in pages:
            table = page.extract_table()
            if table:
                df = pd.DataFrame(table[1:], columns=table[0])
                data.append(df[['Description/Item', 'UOM']])
        return pd.concat(data, ignore_index=True)

# Function to export data to Word, Excel, and text files
def export_data(data):
    # Export to Word
    doc = Document()
    doc.add_heading('PDF Data', level=1)
    for index, row in data.iterrows():
        doc.add_paragraph(f"{row['Description/Item']}, {row['UOM']}")
    doc.save('pdf_data.docx')

    # Export to Excel
    wb = Workbook()
    ws = wb.active
    ws.append(['Description/Item', 'UOM'])
    for index, row in data.iterrows():
        ws.append([row['Description/Item'], row['UOM']])
    wb.save('pdf_data.xlsx')

    # Export to text file
    with open('pdf_data.txt', 'w') as f:
        f.write(data.to_string(index=False))

# Main function
def main():
    pdf_files = [file for file in os.listdir(pdf_directory) if file.endswith('.pdf')]
    all_data = []
    for file in pdf_files:
        pdf_file_path = os.path.join(pdf_directory, file)
        data = extract_data_from_pdf(pdf_file_path)
        all_data.append(data)
    
    all_data_combined = pd.concat(all_data, ignore_index=True)
    export_data(all_data_combined)

if __name__ == "__main__":
    main()
